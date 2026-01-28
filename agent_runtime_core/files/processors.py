"""
Built-in file processors for common file types.

Each processor handles specific file types and extracts text/metadata.
Optional dependencies are checked at runtime.
"""

import csv
import io
import json
import time
from pathlib import Path
from typing import Optional

from .base import (
    FileProcessor,
    FileType,
    ProcessedFile,
    ProcessingOptions,
)


class TextFileProcessor(FileProcessor):
    """Processor for plain text files."""
    
    @property
    def name(self) -> str:
        return "text"
    
    @property
    def supported_types(self) -> list[FileType]:
        return [FileType.TEXT, FileType.MARKDOWN, FileType.JSON, FileType.HTML]
    
    @property
    def supported_extensions(self) -> list[str]:
        return [".txt", ".text", ".log", ".md", ".markdown", ".json", ".html", ".htm", ".xml", ".yaml", ".yml"]
    
    @property
    def supported_mime_types(self) -> list[str]:
        return [
            "text/plain",
            "text/markdown",
            "text/html",
            "application/json",
            "text/xml",
            "application/xml",
            "text/yaml",
        ]
    
    async def process(
        self,
        content: bytes,
        filename: str,
        options: ProcessingOptions,
    ) -> ProcessedFile:
        start_time = time.time()
        
        # Detect encoding
        text = self._decode_text(content)
        
        # Determine specific type
        ext = Path(filename).suffix.lower()
        if ext in [".md", ".markdown"]:
            file_type = FileType.MARKDOWN
        elif ext == ".json":
            file_type = FileType.JSON
        elif ext in [".html", ".htm"]:
            file_type = FileType.HTML
        else:
            file_type = FileType.TEXT
        
        metadata = {
            "line_count": text.count("\n") + 1,
            "char_count": len(text),
            "word_count": len(text.split()),
        }
        
        # For JSON, try to parse and add structure info
        if file_type == FileType.JSON:
            try:
                parsed = json.loads(text)
                metadata["json_type"] = type(parsed).__name__
                if isinstance(parsed, dict):
                    metadata["json_keys"] = list(parsed.keys())[:20]
                elif isinstance(parsed, list):
                    metadata["json_length"] = len(parsed)
            except json.JSONDecodeError:
                pass
        
        return ProcessedFile(
            filename=filename,
            file_type=file_type,
            mime_type=self._get_mime_type(filename),
            size_bytes=len(content),
            text=text,
            metadata=metadata,
            processor_used=self.name,
            processing_time_ms=(time.time() - start_time) * 1000,
        )
    
    def _decode_text(self, content: bytes) -> str:
        """Decode bytes to string, trying multiple encodings."""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        # Fallback with replacement
        return content.decode("utf-8", errors="replace")
    
    def _get_mime_type(self, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        mime_map = {
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".json": "application/json",
            ".html": "text/html",
            ".xml": "application/xml",
            ".yaml": "text/yaml",
        }
        return mime_map.get(ext, "text/plain")


class CsvProcessor(FileProcessor):
    """Processor for CSV files."""
    
    @property
    def name(self) -> str:
        return "csv"
    
    @property
    def supported_types(self) -> list[FileType]:
        return [FileType.CSV]
    
    @property
    def supported_extensions(self) -> list[str]:
        return [".csv", ".tsv"]
    
    @property
    def supported_mime_types(self) -> list[str]:
        return ["text/csv", "text/tab-separated-values"]
    
    async def process(
        self,
        content: bytes,
        filename: str,
        options: ProcessingOptions,
    ) -> ProcessedFile:
        start_time = time.time()
        
        # Decode content
        text = content.decode("utf-8", errors="replace")
        
        # Parse CSV
        delimiter = "\t" if filename.endswith(".tsv") else ","
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = list(reader)
        
        # Extract metadata
        headers = rows[0] if rows else []
        row_count = len(rows) - 1 if rows else 0  # Exclude header

        metadata = {
            "headers": headers,
            "row_count": row_count,
            "column_count": len(headers),
            "delimiter": delimiter,
        }

        return ProcessedFile(
            filename=filename,
            file_type=FileType.CSV,
            mime_type="text/csv",
            size_bytes=len(content),
            text=text,
            metadata=metadata,
            processor_used=self.name,
            processing_time_ms=(time.time() - start_time) * 1000,
        )


class PDFProcessor(FileProcessor):
    """
    Processor for PDF files.

    Requires: pypdf (pip install pypdf)
    """

    @property
    def name(self) -> str:
        return "pdf"

    @property
    def supported_types(self) -> list[FileType]:
        return [FileType.PDF]

    @property
    def supported_extensions(self) -> list[str]:
        return [".pdf"]

    @property
    def supported_mime_types(self) -> list[str]:
        return ["application/pdf"]

    async def process(
        self,
        content: bytes,
        filename: str,
        options: ProcessingOptions,
    ) -> ProcessedFile:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF processing. Install with: pip install pypdf")

        start_time = time.time()
        warnings = []

        # Read PDF
        reader = PdfReader(io.BytesIO(content))

        # Extract text from pages
        text_parts = []
        page_limit = options.pdf_page_limit or len(reader.pages)

        for i, page in enumerate(reader.pages[:page_limit]):
            try:
                text_parts.append(page.extract_text() or "")
            except Exception as e:
                warnings.append(f"Failed to extract text from page {i+1}: {e}")

        text = "\n\n".join(text_parts)

        # Extract metadata
        info = reader.metadata or {}
        metadata = {
            "page_count": len(reader.pages),
            "pages_processed": min(page_limit, len(reader.pages)),
            "title": info.get("/Title", ""),
            "author": info.get("/Author", ""),
            "subject": info.get("/Subject", ""),
            "creator": info.get("/Creator", ""),
            "producer": info.get("/Producer", ""),
        }

        return ProcessedFile(
            filename=filename,
            file_type=FileType.PDF,
            mime_type="application/pdf",
            size_bytes=len(content),
            text=text,
            metadata=metadata,
            processor_used=self.name,
            processing_time_ms=(time.time() - start_time) * 1000,
            warnings=warnings,
        )


class ImageProcessor(FileProcessor):
    """
    Processor for image files.

    Requires: Pillow (pip install Pillow)
    """

    @property
    def name(self) -> str:
        return "image"

    @property
    def supported_types(self) -> list[FileType]:
        return [FileType.IMAGE]

    @property
    def supported_extensions(self) -> list[str]:
        return [".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"]

    @property
    def supported_mime_types(self) -> list[str]:
        return [
            "image/png",
            "image/jpeg",
            "image/gif",
            "image/webp",
            "image/bmp",
            "image/tiff",
        ]

    async def process(
        self,
        content: bytes,
        filename: str,
        options: ProcessingOptions,
    ) -> ProcessedFile:
        try:
            from PIL import Image
        except ImportError:
            raise ImportError("Pillow is required for image processing. Install with: pip install Pillow")

        import base64
        start_time = time.time()

        # Open image
        img = Image.open(io.BytesIO(content))

        # Extract metadata
        metadata = {
            "width": img.width,
            "height": img.height,
            "format": img.format,
            "mode": img.mode,
        }

        # Add EXIF data if available
        if hasattr(img, "_getexif") and img._getexif():
            exif = img._getexif()
            metadata["has_exif"] = True

        # Generate thumbnail
        thumbnail_base64 = None
        if options.generate_thumbnail:
            thumb = img.copy()
            thumb.thumbnail(options.thumbnail_size)
            thumb_buffer = io.BytesIO()
            thumb.save(thumb_buffer, format="PNG")
            thumbnail_base64 = base64.b64encode(thumb_buffer.getvalue()).decode("utf-8")

        # Images don't have text by default - OCR or vision needed
        text = ""

        return ProcessedFile(
            filename=filename,
            file_type=FileType.IMAGE,
            mime_type=f"image/{(img.format or 'png').lower()}",
            size_bytes=len(content),
            text=text,
            metadata=metadata,
            thumbnail_base64=thumbnail_base64,
            processor_used=self.name,
            processing_time_ms=(time.time() - start_time) * 1000,
        )


class DocxProcessor(FileProcessor):
    """
    Processor for Microsoft Word documents.

    Requires: python-docx (pip install python-docx)
    """

    @property
    def name(self) -> str:
        return "docx"

    @property
    def supported_types(self) -> list[FileType]:
        return [FileType.DOCX]

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    @property
    def supported_mime_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]

    async def process(
        self,
        content: bytes,
        filename: str,
        options: ProcessingOptions,
    ) -> ProcessedFile:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

        start_time = time.time()

        # Read document
        doc = docx.Document(io.BytesIO(content))

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)

        if table_texts:
            text += "\n\n--- Tables ---\n" + "\n".join(table_texts)

        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }

        return ProcessedFile(
            filename=filename,
            file_type=FileType.DOCX,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size_bytes=len(content),
            text=text,
            metadata=metadata,
            processor_used=self.name,
            processing_time_ms=(time.time() - start_time) * 1000,
        )


class XlsxProcessor(FileProcessor):
    """
    Processor for Microsoft Excel spreadsheets.

    Requires: openpyxl (pip install openpyxl)
    """

    @property
    def name(self) -> str:
        return "xlsx"

    @property
    def supported_types(self) -> list[FileType]:
        return [FileType.XLSX]

    @property
    def supported_extensions(self) -> list[str]:
        return [".xlsx"]

    @property
    def supported_mime_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]

    async def process(
        self,
        content: bytes,
        filename: str,
        options: ProcessingOptions,
    ) -> ProcessedFile:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for XLSX processing. Install with: pip install openpyxl")

        start_time = time.time()

        # Read workbook
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)

        # Extract text from all sheets
        text_parts = []
        sheet_info = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")

            rows = []
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip(" |"):
                    rows.append(row_text)
                    row_count += 1

            text_parts.extend(rows)
            sheet_info.append({
                "name": sheet_name,
                "row_count": row_count,
            })

        text = "\n".join(text_parts)

        metadata = {
            "sheet_count": len(wb.sheetnames),
            "sheet_names": wb.sheetnames,
            "sheets": sheet_info,
        }

        wb.close()

        return ProcessedFile(
            filename=filename,
            file_type=FileType.XLSX,
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            size_bytes=len(content),
            text=text,
            metadata=metadata,
            processor_used=self.name,
            processing_time_ms=(time.time() - start_time) * 1000,
        )

